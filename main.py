import os
import json
import time
from datetime import datetime

import flet as ft

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============================================================
# CONFIGURACIÓN
# ============================================================

APP_NAME = "FENIX - Sistema de Proformas"

COLOR_PRINCIPAL = "#E63946"
COLOR_VERDE = "#2FA572"
COLOR_GRIS = "#4A4D50"
COLOR_FONDO = "#181818"
COLOR_TARJETA = "#252525"
COLOR_FORM = "#2B2B2B"


# ============================================================
# RUTAS
# ============================================================

def get_storage_path():

    ruta = os.environ.get(
        "FLET_APP_STORAGE_DATA"
    )

    if not ruta:
        ruta = os.path.join(
            os.path.expanduser("~"),
            ".fenix_proformas"
        )

    os.makedirs(
        ruta,
        exist_ok=True
    )

    return ruta


BASE_PATH = get_storage_path()

PRODUCTOS_FILE = os.path.join(
    BASE_PATH,
    "productos.json"
)

IMAGENES_PATH = os.path.join(
    BASE_PATH,
    "imagenes"
)

DOCUMENTOS_PATH = os.path.join(
    BASE_PATH,
    "documentos"
)

# Archivos utilizados para el formato de la proforma
ENCABEZADO_FILE = os.path.join(
    BASE_PATH,
    "encabezado.png"
)

PIE_FILE = os.path.join(
    BASE_PATH,
    "pie_pagina.png"
)

os.makedirs(
    IMAGENES_PATH,
    exist_ok=True
)

os.makedirs(
    DOCUMENTOS_PATH,
    exist_ok=True
)


# ============================================================
# GESTOR DE PRODUCTOS
# ============================================================

class GestorProductos:

    def __init__(self):

        self.archivo = PRODUCTOS_FILE

        self.productos = (
            self.cargar_productos()
        )

    def cargar_productos(self):

        if not os.path.exists(
            self.archivo
        ):
            return []

        try:

            with open(
                self.archivo,
                "r",
                encoding="utf-8"
            ) as archivo:

                datos = json.load(
                    archivo
                )

                if isinstance(
                    datos,
                    list
                ):
                    return datos

        except Exception:

            return []

        return []

    def guardar_archivo(self):

        with open(
            self.archivo,
            "w",
            encoding="utf-8"
        ) as archivo:

            json.dump(
                self.productos,
                archivo,
                indent=4,
                ensure_ascii=False
            )

    def guardar_producto(
        self,
        nombre,
        precio,
        imagen=""
    ):

        for producto in self.productos:

            if (
                producto["nombre"].lower()
                == nombre.lower()
            ):

                return producto

        producto = {
            "nombre": nombre,
            "precio": float(precio),
            "imagen": imagen
        }

        self.productos.append(
            producto
        )

        self.guardar_archivo()

        return producto

    def eliminar_producto(
        self,
        nombre
    ):

        self.productos = [
            producto
            for producto in self.productos
            if producto["nombre"] != nombre
        ]

        self.guardar_archivo()


# ============================================================
# GENERADOR WORD
# ============================================================

class GeneradorWord:

    # --------------------------------------------------------
    # FUENTE GENERAL
    # --------------------------------------------------------

    def _configurar_fuente(
        self,
        run,
        size=10,
        bold=False
    ):

        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold

        # Compatibilidad con Word
        rPr = run._element.get_or_add_rPr()

        rFonts = rPr.rFonts

        if rFonts is None:

            rFonts = OxmlElement(
                "w:rFonts"
            )

            rPr.insert(
                0,
                rFonts
            )

        rFonts.set(
            qn("w:ascii"),
            "Times New Roman"
        )

        rFonts.set(
            qn("w:hAnsi"),
            "Times New Roman"
        )

    # --------------------------------------------------------
    # BORDES
    # --------------------------------------------------------

    def _forzar_bordes_tabla(
        self,
        tabla
    ):

        tblPr = tabla._tbl.tblPr

        tblBorders = tblPr.find(
            qn("w:tblBorders")
        )

        if tblBorders is not None:

            tblPr.remove(
                tblBorders
            )

        tblBorders = OxmlElement(
            "w:tblBorders"
        )

        for border_name in [
            "top",
            "left",
            "bottom",
            "right",
            "insideH",
            "insideV"
        ]:

            border = OxmlElement(
                f"w:{border_name}"
            )

            border.set(
                qn("w:val"),
                "single"
            )

            border.set(
                qn("w:sz"),
                "4"
            )

            border.set(
                qn("w:space"),
                "0"
            )

            border.set(
                qn("w:color"),
                "000000"
            )

            tblBorders.append(
                border
            )

        tblPr.append(
            tblBorders
        )

    # --------------------------------------------------------
    # MÁRGENES DE CELDA
    # --------------------------------------------------------

    def _margenes_celda(
        self,
        celda,
        top=80,
        start=90,
        bottom=80,
        end=90
    ):

        tc = celda._tc

        tcPr = tc.get_or_add_tcPr()

        tcMar = tcPr.first_child_found_in(
            "w:tcMar"
        )

        if tcMar is None:

            tcMar = OxmlElement(
                "w:tcMar"
            )

            tcPr.append(
                tcMar
            )

        for margin, valor in [
            ("top", top),
            ("start", start),
            ("bottom", bottom),
            ("end", end)
        ]:

            nodo = tcMar.find(
                qn(f"w:{margin}")
            )

            if nodo is None:

                nodo = OxmlElement(
                    f"w:{margin}"
                )

                tcMar.append(
                    nodo
                )

            nodo.set(
                qn("w:w"),
                str(valor)
            )

            nodo.set(
                qn("w:type"),
                "dxa"
            )

    # --------------------------------------------------------
    # ALINEACIÓN VERTICAL
    # --------------------------------------------------------

    def _alinear_celda(
        self,
        celda,
        vertical=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    ):

        celda.vertical_alignment = vertical

    # --------------------------------------------------------
    # ESCRIBIR DATO
    # --------------------------------------------------------

    def _escribir_dato(
        self,
        parrafo,
        etiqueta,
        valor,
        size=10,
        centrado=False
    ):

        if centrado:

            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        else:

            parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT

        parrafo.paragraph_format.space_after = Pt(3)
        parrafo.paragraph_format.space_before = Pt(0)
        parrafo.paragraph_format.line_spacing = 1

        r1 = parrafo.add_run(
            etiqueta
        )

        self._configurar_fuente(
            r1,
            size=size,
            bold=True
        )

        r2 = parrafo.add_run(
            str(valor)
        )

        self._configurar_fuente(
            r2,
            size=size,
            bold=False
        )

    # --------------------------------------------------------
    # CONFIGURAR TABLA
    # --------------------------------------------------------

    def _ancho_tabla(
        self,
        tabla,
        ancho_cm=17.5
    ):

        tbl = tabla._tbl

        tblPr = tbl.tblPr

        tblW = tblPr.find(
            qn("w:tblW")
        )

        if tblW is None:

            tblW = OxmlElement(
                "w:tblW"
            )

            tblPr.append(
                tblW
            )

        tblW.set(
            qn("w:w"),
            str(int(ancho_cm * 567))
        )

        tblW.set(
            qn("w:type"),
            "dxa"
        )

    # --------------------------------------------------------
    # GENERAR
    # --------------------------------------------------------

    def generar(
        self,
        datos_form,
        carrito
    ):

        doc = Document()

        # ====================================================
        # CONFIGURACIÓN DE PÁGINA
        # ====================================================

        seccion = doc.sections[0]

        seccion.left_margin = Cm(2.0)
        seccion.right_margin = Cm(2.0)
        seccion.top_margin = Cm(1.5)
        seccion.bottom_margin = Cm(1.4)

        # ====================================================
        # ESTILO GENERAL
        # ====================================================

        estilo_normal = doc.styles["Normal"]

        estilo_normal.font.name = (
            "Times New Roman"
        )

        estilo_normal.font.size = Pt(
            10
        )

        # ====================================================
        # ENCABEZADO
        # ====================================================

        encabezado = seccion.header

        p_enc = encabezado.paragraphs[0]

        p_enc.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        p_enc.paragraph_format.space_after = Pt(0)
        p_enc.paragraph_format.space_before = Pt(0)

        if os.path.exists(
            ENCABEZADO_FILE
        ):

            run = p_enc.add_run()

            run.add_picture(
                ENCABEZADO_FILE,
                width=Cm(17.5)
            )

        else:

            run = p_enc.add_run(
                "FENIX"
            )

            self._configurar_fuente(
                run,
                size=22,
                bold=True
            )

        # ====================================================
        # PIE DE PÁGINA
        # ====================================================

        pie = seccion.footer

        p_pie = pie.paragraphs[0]

        p_pie.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        p_pie.paragraph_format.space_before = Pt(0)
        p_pie.paragraph_format.space_after = Pt(0)

        if os.path.exists(
            PIE_FILE
        ):

            run = p_pie.add_run()

            run.add_picture(
                PIE_FILE,
                width=Cm(14.5)
            )

        else:

            run = p_pie.add_run(
                "Facebook: Fénix systems       "
                "Instagram: fenixsystemsialarms"
            )

            self._configurar_fuente(
                run,
                size=8
            )

        # ====================================================
        # ESPACIO SUPERIOR
        # ====================================================

        p = doc.add_paragraph()

        p.paragraph_format.space_after = Pt(7)

        # ====================================================
        # TITULO
        # ====================================================

        tabla_titulo = doc.add_table(
            rows=1,
            cols=1
        )

        self._forzar_bordes_tabla(
            tabla_titulo
        )

        self._ancho_tabla(
            tabla_titulo
        )

        celda_titulo = (
            tabla_titulo.cell(0, 0)
        )

        self._margenes_celda(
            celda_titulo,
            top=55,
            bottom=55
        )

        celda_titulo.vertical_alignment = (
            WD_CELL_VERTICAL_ALIGNMENT.CENTER
        )

        p_titulo = (
            celda_titulo
            .paragraphs[0]
        )

        p_titulo.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        p_titulo.paragraph_format.space_after = Pt(0)

        run = p_titulo.add_run(
            datos_form["titulo"]
        )

        self._configurar_fuente(
            run,
            size=10,
            bold=True
        )

        doc.add_paragraph()

        # ====================================================
        # DATOS DE EMPRESA / CLIENTE
        # ====================================================

        tabla_datos = doc.add_table(
            rows=1,
            cols=2
        )

        self._forzar_bordes_tabla(
            tabla_datos
        )

        self._ancho_tabla(
            tabla_datos
        )

        # ----------------------------------------------------
        # EMPRESA
        # ----------------------------------------------------

        c1 = tabla_datos.cell(
            0,
            0
        )

        self._margenes_celda(
            c1,
            top=80,
            bottom=80
        )

        c1.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        titulo_empresa = (
            c1.paragraphs[0]
            .add_run(
                "Datos de empresa"
            )
        )

        self._configurar_fuente(
            titulo_empresa,
            size=10,
            bold=True
        )

        self._escribir_dato(
            c1.add_paragraph(),
            "Técnico: ",
            "Teodoro López"
        )

        self._escribir_dato(
            c1.add_paragraph(),
            "Ruc: ",
            "1004320337001"
        )

        self._escribir_dato(
            c1.add_paragraph(),
            "Dirección: ",
            "Ibarra - Ecuador."
        )

        self._escribir_dato(
            c1.add_paragraph(),
            "Teléfono: ",
            "0988195938 / 0992318470"
        )

        self._escribir_dato(
            c1.add_paragraph(),
            "Correo: ",
            "teo.master@hotmail.com"
        )

        self._escribir_dato(
            c1.add_paragraph(),
            "Obligado a llevar contabilidad: ",
            "No"
        )

        # ----------------------------------------------------
        # CLIENTE
        # ----------------------------------------------------

        c2 = tabla_datos.cell(
            0,
            1
        )

        self._margenes_celda(
            c2,
            top=80,
            bottom=80
        )

        c2.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        titulo_cliente = (
            c2.paragraphs[0]
            .add_run(
                "Datos del cliente"
            )
        )

        self._configurar_fuente(
            titulo_cliente,
            size=10,
            bold=True
        )

        self._escribir_dato(
            c2.add_paragraph(),
            "Destinatario: ",
            datos_form["destinatario"]
        )

        self._escribir_dato(
            c2.add_paragraph(),
            "Dirección: ",
            datos_form["direccion"]
        )

        self._escribir_dato(
            c2.add_paragraph(),
            "Teléfono: ",
            datos_form["telefono"]
        )

        self._escribir_dato(
            c2.add_paragraph(),
            "Correo: ",
            datos_form["correo"]
        )

        doc.add_paragraph()

        # ====================================================
        # FECHA / DATOS ADICIONALES
        # ====================================================

        tabla_extras = doc.add_table(
            rows=1,
            cols=2
        )

        self._forzar_bordes_tabla(
            tabla_extras
        )

        self._ancho_tabla(
            tabla_extras
        )

        # ----------------------------------------------------
        # IZQUIERDA
        # ----------------------------------------------------

        c_ext1 = tabla_extras.cell(
            0,
            0
        )

        self._margenes_celda(
            c_ext1,
            top=80,
            bottom=80
        )

        fecha_hoy = datetime.now().strftime(
            "%d-%m-%Y"
        )

        self._escribir_dato(
            c_ext1.paragraphs[0],
            "Fecha de emisión: ",
            fecha_hoy
        )

        self._escribir_dato(
            c_ext1.add_paragraph(),
            "Vigencia de oferta: ",
            datos_form["vigencia"]
        )

        # ----------------------------------------------------
        # DERECHA
        # ----------------------------------------------------

        c_ext2 = tabla_extras.cell(
            0,
            1
        )

        self._margenes_celda(
            c_ext2,
            top=80,
            bottom=80
        )

        c_ext2.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        titulo_extra = (
            c_ext2.paragraphs[0]
            .add_run(
                "Datos adicionales"
            )
        )

        self._configurar_fuente(
            titulo_extra,
            size=10,
            bold=True
        )

        self._escribir_dato(
            c_ext2.add_paragraph(),
            "Concepto: ",
            datos_form["concepto"]
        )

        self._escribir_dato(
            c_ext2.add_paragraph(),
            "Forma de pago: ",
            datos_form["pago"]
        )

        doc.add_paragraph()

        # ====================================================
        # TABLA DE PRODUCTOS
        # ====================================================

        tabla_items = doc.add_table(
            rows=1,
            cols=4
        )

        self._forzar_bordes_tabla(
            tabla_items
        )

        self._ancho_tabla(
            tabla_items
        )

        encabezados = [
            "Cant.",
            "Descripción",
            "V. Unitario",
            "V. Total"
        ]

        for i, texto in enumerate(
            encabezados
        ):

            celda = (
                tabla_items
                .rows[0]
                .cells[i]
            )

            self._margenes_celda(
                celda,
                top=55,
                bottom=55
            )

            p = celda.paragraphs[0]

            p.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
            )

            run = p.add_run(
                texto
            )

            self._configurar_fuente(
                run,
                size=9,
                bold=True
            )

        # ====================================================
        # FILAS PRODUCTOS
        # ====================================================

        total = 0.0

        for descripcion, datos in (
            carrito.items()
        ):

            cantidad = int(
                datos["cant"]
            )

            precio = float(
                datos["precio"]
            )

            subtotal = (
                cantidad * precio
            )

            total += subtotal

            fila = (
                tabla_items
                .add_row()
                .cells
            )

            valores = [
                str(cantidad),
                str(descripcion),
                f"${precio:.2f}",
                f"${subtotal:.2f}"
            ]

            for i, valor in enumerate(
                valores
            ):

                self._margenes_celda(
                    fila[i],
                    top=45,
                    bottom=45
                )

                p = fila[i].paragraphs[0]

                p.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                )

                run = p.add_run(
                    valor
                )

                self._configurar_fuente(
                    run,
                    size=9
                )

        # ====================================================
        # TOTAL
        # ====================================================

        p_total = doc.add_paragraph()

        p_total.alignment = (
            WD_ALIGN_PARAGRAPH.RIGHT
        )

        p_total.paragraph_format.space_before = Pt(0)
        p_total.paragraph_format.space_after = Pt(3)

        run_total = p_total.add_run(
            f"TOTAL A PAGAR: ${total:.2f}"
        )

        self._configurar_fuente(
            run_total,
            size=10,
            bold=True
        )

        # ====================================================
        # GARANTÍA
        # ====================================================

        doc.add_paragraph()

        tabla_garantia = doc.add_table(
            rows=1,
            cols=1
        )

        self._forzar_bordes_tabla(
            tabla_garantia
        )

        self._ancho_tabla(
            tabla_garantia
        )

        celda_garantia = (
            tabla_garantia.cell(
                0,
                0
            )
        )

        self._margenes_celda(
            celda_garantia,
            top=80,
            bottom=80,
            start=90,
            end=90
        )

        p_garantia = (
            celda_garantia
            .paragraphs[0]
        )

        p_garantia.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT
        )

        p_garantia.paragraph_format.line_spacing = 1

        texto_garantia = (
            "La garantía de los equipos será de un año, "
            "solo y bajo defectos de fábrica, Fénix "
            "sistemas de seguridad no se responsabiliza "
            "de daños por falta de mantenimiento o "
            "cuidados (responsabilidad del cliente). "
            "Para más información consulte el catálogo "
            "de productos."
        )

        run_garantia = (
            p_garantia.add_run(
                texto_garantia
            )
        )

        self._configurar_fuente(
            run_garantia,
            size=9
        )

        # ====================================================
        # GUARDAR
        # ====================================================

        codigo_unico = int(
            time.time()
        )

        nombre_limpio = (
            datos_form["destinatario"]
            .strip()
            .replace(" ", "_")
        )

        if not nombre_limpio:

            nombre_limpio = (
                "Cotizacion"
            )

        ruta_documento = os.path.join(
            DOCUMENTOS_PATH,
            f"Proforma_"
            f"{nombre_limpio}_"
            f"{codigo_unico}.docx"
        )

        doc.save(
            ruta_documento
        )

        return ruta_documento


# ============================================================
# APLICACIÓN
# ============================================================

class AplicacionProformas:

    def __init__(
        self,
        page: ft.Page
    ):

        self.page = page

        # ----------------------------------------------------
        # DATOS
        # ----------------------------------------------------

        self.db_productos = (
            GestorProductos()
        )

        self.generador_word = (
            GeneradorWord()
        )

        self.carrito = {}

        self.filas_manuales = []

        # ----------------------------------------------------
        # CONFIGURACIÓN
        # ----------------------------------------------------

        self.page.title = APP_NAME
        self.page.bgcolor = COLOR_FONDO
        self.page.padding = 0
        self.page.scroll = ft.ScrollMode.AUTO
        self.page.theme_mode = ft.ThemeMode.DARK

        self.page.theme = ft.Theme(
            color_scheme_seed=COLOR_PRINCIPAL
        )

        # ----------------------------------------------------
        # FILE PICKER
        # ----------------------------------------------------

        self.file_picker = ft.FilePicker()

        self.page.services.append(
            self.file_picker
        )

        self.imagen_bytes = None
        self.imagen_nombre = ""

        # ----------------------------------------------------
        # CONTENEDORES
        # ----------------------------------------------------

        self.catalogo_view = ft.Column(
            spacing=10
        )

        self.carrito_view = ft.Column(
            spacing=7
        )

        self.manual_view = ft.Column(
            spacing=8
        )

        # ----------------------------------------------------
        # CAMPOS
        # ----------------------------------------------------

        self.ent_titulo = ft.TextField(
            label="Título",
            value="COTIZACIÓN ALARMAS :000078539",
            expand=True
        )

        self.ent_dest = ft.TextField(
            label="Destinatario",
            expand=True
        )

        self.ent_dir = ft.TextField(
            label="Dirección",
            expand=True
        )

        self.ent_tel = ft.TextField(
            label="Teléfono",
            expand=True
        )

        self.ent_correo = ft.TextField(
            label="Correo",
            expand=True
        )

        self.ent_concepto = ft.TextField(
            label="Concepto",
            expand=True
        )

        self.ent_pago = ft.TextField(
            label="Forma de pago",
            value="50% por anticipado",
            expand=True
        )

        self.ent_vigencia = ft.TextField(
            label="Vigencia",
            value="7 días",
            expand=True
        )

        self.lbl_total = ft.Text(
            "TOTAL: $0.00",
            size=25,
            weight=ft.FontWeight.BOLD,
            color=COLOR_VERDE
        )

        self.mensaje = ft.Text(
            ""
        )

        # ----------------------------------------------------
        # INTERFAZ
        # ----------------------------------------------------

        self.construir_interfaz()

        self.agregar_fila_manual(
            actualizar=False
        )

        self.dibujar_catalogo()

        self.actualizar_carrito()

    # ========================================================
    # INTERFAZ
    # ========================================================

    def crear_panel(
        self,
        titulo,
        contenido,
        color=COLOR_TARJETA
    ):

        return ft.Container(
            bgcolor=color,
            border_radius=14,
            padding=14,
            content=ft.Column(
                [
                    ft.Text(
                        titulo,
                        size=19,
                        weight=ft.FontWeight.BOLD
                    ),

                    contenido
                ],
                spacing=10
            )
        )

    def construir_interfaz(
        self
    ):

        encabezado = ft.Container(
            bgcolor="#111111",
            padding=12,
            content=ft.Row(
                [
                    ft.Text(
                        "FENIX",
                        size=28,
                        weight=ft.FontWeight.BOLD,
                        color=COLOR_PRINCIPAL
                    ),

                    ft.Text(
                        "Sistema de Proformas",
                        size=16
                    )
                ]
            )
        )

        boton_limpiar = ft.Button(
            content="Limpiar Todo",
            icon=ft.Icons.DELETE_SWEEP,
            on_click=self.limpiar_todo,
            style=ft.ButtonStyle(
                bgcolor=COLOR_GRIS,
                color="#FFFFFF"
            )
        )

        boton_nuevo = ft.Button(
            content="Añadir Nuevo Producto",
            icon=ft.Icons.ADD_BOX,
            on_click=self.abrir_nuevo_producto,
            style=ft.ButtonStyle(
                bgcolor=COLOR_VERDE,
                color="#FFFFFF"
            )
        )

        panel_catalogo = self.crear_panel(
            "Catálogo de Productos",
            ft.Column(
                [
                    ft.Row(
                        [
                            boton_limpiar,
                            boton_nuevo
                        ],
                        wrap=True,
                        spacing=8
                    ),

                    self.catalogo_view
                ]
            )
        )

        formulario = ft.Column(
            [
                self.ent_titulo,

                ft.Row(
                    [
                        self.ent_dest,
                        self.ent_dir
                    ],
                    spacing=8
                ),

                ft.Row(
                    [
                        self.ent_tel,
                        self.ent_correo
                    ],
                    spacing=8
                ),

                self.ent_concepto,

                ft.Row(
                    [
                        self.ent_pago,
                        self.ent_vigencia
                    ],
                    spacing=8
                )
            ],
            spacing=8
        )

        panel_formulario = self.crear_panel(
            "Datos del Documento",
            formulario,
            COLOR_FORM
        )

        boton_agregar_fila = ft.Button(
            content="Agregar producto",
            icon=ft.Icons.ADD,
            on_click=self.agregar_fila_manual,
            style=ft.ButtonStyle(
                bgcolor=COLOR_VERDE,
                color="#FFFFFF"
            )
        )

        panel_manual = self.crear_panel(
            "Añadir Manualmente",
            ft.Column(
                [
                    self.manual_view,

                    ft.Row(
                        [
                            boton_agregar_fila
                        ],
                        alignment=(
                            ft.MainAxisAlignment.END
                        )
                    )
                ]
            ),
            COLOR_FORM
        )

        panel_carrito = self.crear_panel(
            "Carrito",
            ft.Column(
                [
                    self.carrito_view,

                    ft.Divider(),

                    ft.Row(
                        [
                            self.lbl_total
                        ],
                        alignment=(
                            ft.MainAxisAlignment.END
                        )
                    )
                ]
            )
        )

        boton_generar = ft.Button(
            content="GENERAR WORD",
            icon=ft.Icons.DESCRIPTION,
            on_click=self.ejecutar_generacion,
            style=ft.ButtonStyle(
                bgcolor=COLOR_PRINCIPAL,
                color="#FFFFFF",
                padding=18
            )
        )

        panel_derecho = ft.Column(
            [
                panel_formulario,
                panel_manual,
                panel_carrito,

                ft.Row(
                    [
                        boton_generar
                    ],
                    alignment=(
                        ft.MainAxisAlignment.CENTER
                    )
                ),

                self.mensaje
            ],
            spacing=12
        )

        layout = ft.ResponsiveRow(
            [
                ft.Container(
                    content=panel_catalogo,
                    col={
                        "xs": 12,
                        "sm": 12,
                        "md": 6,
                        "lg": 6
                    }
                ),

                ft.Container(
                    content=panel_derecho,
                    col={
                        "xs": 12,
                        "sm": 12,
                        "md": 6,
                        "lg": 6
                    }
                )
            ],
            spacing=12
        )

        self.page.add(
            ft.Column(
                [
                    encabezado,

                    ft.Container(
                        padding=12,
                        content=layout
                    )
                ],
                spacing=0
            )
        )

    # ========================================================
    # CATALOGO
    # ========================================================

    def dibujar_catalogo(
        self
    ):

        self.catalogo_view.controls.clear()

        productos = (
            self.db_productos.productos
        )

        if not productos:

            self.catalogo_view.controls.append(
                ft.Container(
                    padding=30,
                    alignment=ft.Alignment.CENTER,
                    content=ft.Text(
                        "No hay productos registrados.",
                        color="#999999"
                    )
                )
            )

            self.page.update()

            return

        controles = []

        for producto in productos:

            tarjeta = (
                self.crear_tarjeta_producto(
                    producto
                )
            )

            controles.append(
                ft.Container(
                    content=tarjeta,
                    col={
                        "xs": 6,
                        "sm": 6,
                        "md": 6
                    }
                )
            )

        self.catalogo_view.controls.append(
            ft.ResponsiveRow(
                controles,
                spacing=8
            )
        )

        self.page.update()

    # ========================================================
    # TARJETA PRODUCTO
    # ========================================================

    def crear_tarjeta_producto(
        self,
        producto
    ):

        nombre = producto["nombre"]

        precio = float(
            producto["precio"]
        )

        cantidad = int(
            self.carrito
            .get(
                nombre,
                {}
            )
            .get(
                "cant",
                0
            )
        )

        imagen = self.crear_imagen(
            producto
        )

        texto_cantidad = ft.Text(
            str(cantidad),
            size=18,
            weight=ft.FontWeight.BOLD,
            width=30,
            text_align=ft.TextAlign.CENTER
        )

        boton_menos = ft.IconButton(
            icon=ft.Icons.REMOVE,
            bgcolor=COLOR_GRIS,
            icon_color="#FFFFFF",
            tooltip="Restar",

            on_click=lambda e,
            n=nombre,
            p=precio:
                self.modificar_carrito(
                    n,
                    p,
                    -1
                )
        )

        boton_mas = ft.IconButton(
            icon=ft.Icons.ADD,
            bgcolor=COLOR_VERDE,
            icon_color="#FFFFFF",
            tooltip="Agregar",

            on_click=lambda e,
            n=nombre,
            p=precio:
                self.modificar_carrito(
                    n,
                    p,
                    1
                )
        )

        controles = ft.Row(
            [
                boton_menos,
                texto_cantidad,
                boton_mas
            ],
            alignment=(
                ft.MainAxisAlignment.CENTER
            ),
            spacing=2
        )

        boton_borrar = ft.TextButton(
            content="Borrar",
            icon=ft.Icons.DELETE,
            icon_color=COLOR_PRINCIPAL,

            on_click=lambda e,
            n=nombre:
                self.eliminar_de_bd(
                    n
                )
        )

        return ft.Container(
            bgcolor="#292929",
            border_radius=12,
            padding=8,

            content=ft.Column(
                [
                    imagen,

                    ft.Text(
                        nombre,
                        size=14,
                        weight=ft.FontWeight.BOLD,
                        max_lines=2,
                        text_align=ft.TextAlign.CENTER
                    ),

                    ft.Text(
                        f"${precio:.2f}",
                        size=14,
                        color="#BBBBBB"
                    ),

                    controles,

                    boton_borrar
                ],

                horizontal_alignment=(
                    ft.CrossAxisAlignment.CENTER
                ),

                spacing=4
            )
        )

    # ========================================================
    # IMAGEN PRODUCTO
    # ========================================================

    def crear_imagen(
        self,
        producto
    ):

        ruta = producto.get(
            "imagen",
            ""
        )

        if ruta and os.path.exists(
            ruta
        ):

            return ft.Container(
                width=110,
                height=90,
                border_radius=8,
                clip_behavior=(
                    ft.ClipBehavior.ANTI_ALIAS
                ),
                content=ft.Image(
                    src=ruta,
                    width=110,
                    height=90,
                    fit=ft.BoxFit.COVER
                )
            )

        return ft.Container(
            width=110,
            height=90,
            bgcolor="#333333",
            border_radius=8,
            alignment=ft.Alignment.CENTER,
            content=ft.Column(
                [
                    ft.Icon(
                        ft.Icons.IMAGE_NOT_SUPPORTED,
                        size=30,
                        color="#777777"
                    ),

                    ft.Text(
                        "Sin Foto",
                        size=12,
                        color="#AAAAAA"
                    )
                ],

                horizontal_alignment=(
                    ft.CrossAxisAlignment.CENTER
                ),

                alignment=(
                    ft.MainAxisAlignment.CENTER
                )
            )
        )

    # ========================================================
    # MODIFICAR CARRITO
    # ========================================================

    def modificar_carrito(
        self,
        descripcion,
        precio,
        cantidad
    ):

        if descripcion in self.carrito:

            self.carrito[
                descripcion
            ]["cant"] += cantidad

        else:

            if cantidad > 0:

                self.carrito[
                    descripcion
                ] = {
                    "cant": cantidad,
                    "precio": float(precio)
                }

        if (
            descripcion in self.carrito
            and self.carrito[
                descripcion
            ]["cant"] <= 0
        ):

            del self.carrito[
                descripcion
            ]

        self.actualizar_carrito()

        self.dibujar_catalogo()

    # ========================================================
    # CARRITO
    # ========================================================

    def actualizar_carrito(
        self
    ):

        self.carrito_view.controls.clear()

        total = 0.0

        for descripcion, datos in (
            self.carrito.items()
        ):

            cantidad = int(
                datos["cant"]
            )

            precio = float(
                datos["precio"]
            )

            subtotal = (
                cantidad * precio
            )

            total += subtotal

            boton_menos = ft.IconButton(
                icon=ft.Icons.REMOVE,
                icon_size=18,

                on_click=lambda e,
                d=descripcion,
                p=precio:
                    self.modificar_carrito(
                        d,
                        p,
                        -1
                    )
            )

            boton_mas = ft.IconButton(
                icon=ft.Icons.ADD,
                icon_size=18,
                bgcolor=COLOR_VERDE,
                icon_color="#FFFFFF",

                on_click=lambda e,
                d=descripcion,
                p=precio:
                    self.modificar_carrito(
                        d,
                        p,
                        1
                    )
            )

            boton_eliminar = ft.IconButton(
                icon=ft.Icons.CLOSE,
                icon_color=COLOR_PRINCIPAL,

                on_click=lambda e,
                d=descripcion:
                    self.eliminar_del_carrito(
                        d
                    )
            )

            fila = ft.Container(
                bgcolor="#202020",
                border_radius=9,
                padding=6,

                content=ft.Row(
                    [
                        ft.Text(
                            descripcion,
                            expand=True,
                            size=14
                        ),

                        boton_menos,

                        ft.Text(
                            str(cantidad),
                            width=25,
                            text_align=(
                                ft.TextAlign.CENTER
                            )
                        ),

                        boton_mas,

                        ft.Text(
                            f"${subtotal:.2f}",
                            width=70,
                            text_align=(
                                ft.TextAlign.RIGHT
                            ),
                            weight=(
                                ft.FontWeight.BOLD
                            )
                        ),

                        boton_eliminar
                    ],

                    spacing=2
                )
            )

            self.carrito_view.controls.append(
                fila
            )

        self.lbl_total.value = (
            f"TOTAL: ${total:.2f}"
        )

        self.page.update()

    # ========================================================
    # ELIMINAR DEL CARRITO
    # ========================================================

    def eliminar_del_carrito(
        self,
        descripcion
    ):

        if descripcion in self.carrito:

            del self.carrito[
                descripcion
            ]

        self.actualizar_carrito()

        self.dibujar_catalogo()

    # ========================================================
    # AÑADIR FILA MANUAL
    # ========================================================

    def agregar_fila_manual(
        self,
        e=None,
        actualizar=True
    ):

        nombre = ft.TextField(
            label="Nombre del producto",
            expand=True
        )

        cantidad = ft.TextField(
            label="Cantidad",
            width=90,
            value="1",
            keyboard_type=(
                ft.KeyboardType.NUMBER
            )
        )

        precio = ft.TextField(
            label="Precio",
            width=110,
            keyboard_type=(
                ft.KeyboardType.NUMBER
            )
        )

        # ----------------------------------------------------
        # AGREGAR ESTA FILA AL CARRITO
        # ----------------------------------------------------

        def agregar_al_carrito(
            e
        ):

            nombre_producto = (
                nombre.value or ""
            ).strip()

            if not nombre_producto:

                self.mostrar_mensaje(
                    "Ingresa el nombre del producto.",
                    COLOR_PRINCIPAL
                )

                return

            try:

                cantidad_num = int(
                    cantidad.value or "0"
                )

                precio_num = float(
                    precio.value or "0"
                )

            except ValueError:

                self.mostrar_mensaje(
                    "Cantidad o precio inválido.",
                    COLOR_PRINCIPAL
                )

                return

            if cantidad_num <= 0:

                self.mostrar_mensaje(
                    "La cantidad debe ser mayor a 0.",
                    COLOR_PRINCIPAL
                )

                return

            if precio_num < 0:

                self.mostrar_mensaje(
                    "El precio no puede ser negativo.",
                    COLOR_PRINCIPAL
                )

                return

            # ----------------------------------------------
            # AQUÍ SE MANDA AL CARRITO
            # ----------------------------------------------

            self.modificar_carrito(
                nombre_producto,
                precio_num,
                cantidad_num
            )

            self.mostrar_mensaje(
                f"{nombre_producto} agregado al carrito.",
                COLOR_VERDE
            )

        # ----------------------------------------------------
        # BOTON AGREGAR
        # ----------------------------------------------------

        boton_agregar = ft.Button(
            content="Agregar al carrito",
            icon=ft.Icons.ADD_SHOPPING_CART,
            on_click=agregar_al_carrito,

            style=ft.ButtonStyle(
                bgcolor=COLOR_VERDE,
                color="#FFFFFF"
            )
        )

        # ----------------------------------------------------
        # BOTON ELIMINAR FILA
        # ----------------------------------------------------

        fila = None

        def eliminar_fila(
            e
        ):

            if fila in self.filas_manuales:

                self.filas_manuales.remove(
                    fila
                )

            if (
                fila
                in self.manual_view.controls
            ):

                self.manual_view.controls.remove(
                    fila
                )

            self.page.update()

        boton_eliminar = ft.IconButton(
            icon=ft.Icons.DELETE,
            icon_color=COLOR_PRINCIPAL,
            tooltip="Eliminar fila",
            on_click=eliminar_fila
        )

        # ----------------------------------------------------
        # FILA
        # ----------------------------------------------------

        fila = ft.Container(
            bgcolor="#202020",
            border_radius=9,
            padding=7,

            content=ft.ResponsiveRow(
                [
                    ft.Container(
                        content=nombre,
                        col={
                            "xs": 12,
                            "sm": 5,
                            "md": 5
                        }
                    ),

                    ft.Container(
                        content=cantidad,
                        col={
                            "xs": 4,
                            "sm": 2,
                            "md": 2
                        }
                    ),

                    ft.Container(
                        content=precio,
                        col={
                            "xs": 4,
                            "sm": 2,
                            "md": 2
                        }
                    ),

                    ft.Container(
                        content=boton_agregar,
                        col={
                            "xs": 3,
                            "sm": 2,
                            "md": 2
                        }
                    ),

                    ft.Container(
                        content=boton_eliminar,
                        col={
                            "xs": 1,
                            "sm": 1,
                            "md": 1
                        }
                    )
                ],
                spacing=5
            )
        )

        self.filas_manuales.append(
            fila
        )

        self.manual_view.controls.append(
            fila
        )

        if actualizar:

            self.page.update()

    # ========================================================
    # NUEVO PRODUCTO
    # ========================================================

    def abrir_nuevo_producto(
        self,
        e
    ):

        self.imagen_bytes = None

        self.imagen_nombre = ""

        self.ent_nuevo_nombre = ft.TextField(
            label="Nombre del producto"
        )

        self.ent_nuevo_precio = ft.TextField(
            label="Precio",
            keyboard_type=(
                ft.KeyboardType.NUMBER
            )
        )

        self.lbl_imagen = ft.Text(
            "Sin imagen"
        )

        self.dialogo_producto = ft.AlertDialog(

            modal=True,

            title=ft.Text(
                "Nuevo Producto"
            ),

            content=ft.Column(
                [
                    self.ent_nuevo_nombre,
                    self.ent_nuevo_precio,
                    self.lbl_imagen
                ],
                tight=True,
                width=350
            ),

            actions=[
                ft.TextButton(
                    content="Cancelar",
                    on_click=self.cerrar_dialogo
                ),

                ft.TextButton(
                    content="Cargar Imagen",
                    icon=ft.Icons.IMAGE,
                    on_click=self.seleccionar_imagen
                ),

                ft.Button(
                    content="Guardar",
                    icon=ft.Icons.SAVE,

                    on_click=self.guardar_nuevo_producto,

                    style=ft.ButtonStyle(
                        bgcolor=COLOR_VERDE,
                        color="#FFFFFF"
                    )
                )
            ]
        )

        self.page.show_dialog(
            self.dialogo_producto
        )

    # ========================================================
    # SELECCIONAR IMAGEN
    # ========================================================

    async def seleccionar_imagen(
        self,
        e
    ):

        archivos = await self.file_picker.pick_files(
            allow_multiple=False,
            with_data=True
        )

        if not archivos:

            return

        archivo = archivos[0]

        self.imagen_nombre = (
            archivo.name
        )

        self.imagen_bytes = (
            archivo.bytes
        )

        self.lbl_imagen.value = (
            f"✅ {archivo.name}"
        )

        self.page.update()

    # ========================================================
    # GUARDAR NUEVO PRODUCTO
    # ========================================================

    async def guardar_nuevo_producto(
        self,
        e
    ):

        nombre = (
            self.ent_nuevo_nombre.value
            or ""
        ).strip()

        precio_texto = (
            self.ent_nuevo_precio.value
            or ""
        ).strip()

        if not nombre:

            self.mostrar_mensaje(
                "Escribe el nombre.",
                COLOR_PRINCIPAL
            )

            return

        try:

            precio = float(
                precio_texto
            )

        except ValueError:

            self.mostrar_mensaje(
                "Precio inválido.",
                COLOR_PRINCIPAL
            )

            return

        for producto in (
            self.db_productos.productos
        ):

            if (
                producto["nombre"].lower()
                == nombre.lower()
            ):

                self.mostrar_mensaje(
                    "Ese producto ya existe.",
                    COLOR_PRINCIPAL
                )

                return

        ruta_imagen = ""

        if self.imagen_bytes:

            extension = os.path.splitext(
                self.imagen_nombre
            )[1].lower()

            if extension not in [
                ".png",
                ".jpg",
                ".jpeg",
                ".webp"
            ]:

                extension = ".jpg"

            nombre_archivo = (
                self.limpiar_nombre(
                    nombre
                )
                + extension
            )

            ruta_imagen = os.path.join(
                IMAGENES_PATH,
                nombre_archivo
            )

            with open(
                ruta_imagen,
                "wb"
            ) as archivo:

                archivo.write(
                    self.imagen_bytes
                )

        self.db_productos.guardar_producto(
            nombre,
            precio,
            ruta_imagen
        )

        self.cerrar_dialogo()

        self.dibujar_catalogo()

        self.mostrar_mensaje(
            "Producto guardado correctamente.",
            COLOR_VERDE
        )

    # ========================================================
    # CERRAR DIALOGO
    # ========================================================

    def cerrar_dialogo(
        self,
        e=None
    ):

        try:

            self.page.close(
                self.dialogo_producto
            )

        except Exception:

            try:

                self.dialogo_producto.open = False

                self.page.update()

            except Exception:

                pass

    # ========================================================
    # ELIMINAR PRODUCTO BD
    # ========================================================

    def eliminar_de_bd(
        self,
        nombre
    ):

        if nombre in self.carrito:

            del self.carrito[
                nombre
            ]

        self.db_productos.eliminar_producto(
            nombre
        )

        self.actualizar_carrito()

        self.dibujar_catalogo()

        self.mostrar_mensaje(
            "Producto eliminado.",
            COLOR_VERDE
        )

    # ========================================================
    # LIMPIAR
    # ========================================================

    def limpiar_todo(
        self,
        e=None
    ):

        self.carrito.clear()

        self.ent_dest.value = ""
        self.ent_dir.value = ""
        self.ent_tel.value = ""
        self.ent_correo.value = ""
        self.ent_concepto.value = ""

        for fila in self.filas_manuales:

            try:

                responsive = fila.content

                controles = []

                for control in responsive.controls:

                    if hasattr(
                        control,
                        "content"
                    ):

                        controles.append(
                            control.content
                        )

                if len(controles) >= 3:

                    controles[0].value = ""
                    controles[1].value = "1"
                    controles[2].value = ""

            except Exception:

                pass

        self.actualizar_carrito()

        self.dibujar_catalogo()

        self.mensaje.value = ""

        self.page.update()

    # ========================================================
    # GENERAR WORD
    # ========================================================

    async def ejecutar_generacion(
        self,
        e
    ):

        if not self.carrito:

            self.mostrar_mensaje(
                "El carrito está vacío.",
                COLOR_PRINCIPAL
            )

            return

        datos = {

            "titulo": (
                self.ent_titulo.value
                or "COTIZACIÓN"
            ),

            "destinatario": (
                self.ent_dest.value
                or "Consumidor Final"
            ),

            "direccion": (
                self.ent_dir.value
                or "A convenir"
            ),

            "telefono": (
                self.ent_tel.value
                or "A convenir"
            ),

            "correo": (
                self.ent_correo.value
                or "A convenir"
            ),

            "concepto": (
                self.ent_concepto.value
                or "N/A"
            ),

            "pago": (
                self.ent_pago.value
                or "A convenir"
            ),

            "vigencia": (
                self.ent_vigencia.value
                or "A convenir"
            )
        }

        try:

            ruta = (
                self.generador_word.generar(
                    datos,
                    self.carrito
                )
            )

            nombre = os.path.basename(
                ruta
            )

            self.mostrar_mensaje(
                f"Documento generado: {nombre}",
                COLOR_VERDE
            )

            await self.guardar_documento(
                ruta,
                nombre
            )

        except Exception as error:

            self.mostrar_mensaje(
                f"Error generando Word: {error}",
                COLOR_PRINCIPAL
            )

    # ========================================================
    # GUARDAR DOCUMENTO
    # ========================================================

    async def guardar_documento(
        self,
        ruta,
        nombre
    ):

        try:

            with open(
                ruta,
                "rb"
            ) as archivo:

                datos = archivo.read()

            await self.file_picker.save_file(
                dialog_title="Guardar Proforma",
                file_name=nombre,
                src_bytes=datos
            )

        except Exception:

            pass

    # ========================================================
    # LIMPIAR NOMBRE
    # ========================================================

    def limpiar_nombre(
        self,
        texto
    ):

        caracteres = (
            '<>:"/\\|?*'
        )

        resultado = ""

        for caracter in texto:

            if caracter in caracteres:

                resultado += "_"

            else:

                resultado += caracter

        resultado = resultado.strip()

        if not resultado:

            resultado = "producto"

        return resultado[:80]

    # ========================================================
    # MENSAJE
    # ========================================================

    def mostrar_mensaje(
        self,
        texto,
        color="#FFFFFF"
    ):

        self.mensaje.value = texto

        self.mensaje.color = color

        self.page.update()


# ============================================================
# MAIN
# ============================================================

def main(
    page: ft.Page
):

    AplicacionProformas(
        page
    )


if __name__ == "__main__":

    ft.run(
        main
    )